import io
import re
import unicodedata
import zipfile
from datetime import datetime

import pandas as pd
import streamlit as st
from pypdf import PdfReader, PdfWriter
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from xml.sax.saxutils import escape


st.set_page_config(
    page_title="InformeFusion",
    page_icon="📄",
    layout="wide",
)


BLOCS = {
    "informe_1": "Informe oficial - part 1",
    "informe_2": "Informe oficial - part 2",
    "psi": "PSI / Pla de suport individualitzat",
    "observacions": "Observacions escrites",
}

ORDRE_PER_DEFECTE = ["informe_1", "informe_2", "psi", "observacions"]
PARAULES_IGNORADES = {
    "pdf", "informe", "informes", "oficial", "notes", "nota", "creuetes", "creueta",
    "psi", "pis", "pla", "suport", "individualitzat", "observacions", "observacio",
    "part", "pagina", "pag", "trim", "trimestre", "final", "alumne", "alumna",
    "educacio", "primaria", "classe", "curs", "esfera",
}


def normalitza_text(text: str) -> str:
    """Converteix un text a una forma fàcil de comparar."""
    if text is None:
        return ""
    text = str(text).strip().lower()
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def tokens_text(text: str, ignora_generiques: bool = False) -> list[str]:
    """Retorna paraules normalitzades útils per comparar noms."""
    text_norm = normalitza_text(text)
    tokens = [token for token in text_norm.split() if token]
    if ignora_generiques:
        tokens = [t for t in tokens if t not in PARAULES_IGNORADES]
    return tokens


def nom_seguretat_fitxer(text: str) -> str:
    """Neteja el nom final perquè sigui vàlid com a nom de fitxer."""
    text = str(text).strip()
    text = re.sub(r"[\\/:*?\"<>|]", "-", text)
    text = re.sub(r"\s+", " ", text)
    return text[:180]


def llegeix_alumnes(uploaded_file, sense_capcalera=False):
    if uploaded_file is None:
        return None

    nom = uploaded_file.name.lower()
    header = None if sense_capcalera else 0

    if nom.endswith(".csv"):
        return pd.read_csv(uploaded_file, header=header)
    return pd.read_excel(uploaded_file, header=header)


def extreu_nom_alumne(row, mode, columna_nom, columna_cognoms, columna_nom_complet):
    if mode == "Una columna amb nom complet":
        return str(row[columna_nom_complet]).strip()
    nom = str(row[columna_nom]).strip()
    cognoms = str(row[columna_cognoms]).strip()
    return f"{cognoms}, {nom}".strip(", ")


def separa_nom_cognoms(nom_alumne: str):
    """Intenta separar cognoms i nom quan el format és 'Cognoms, Nom'."""
    if "," in str(nom_alumne):
        cognoms, nom = str(nom_alumne).split(",", 1)
        return tokens_text(nom, True), tokens_text(cognoms, True)
    tokens = tokens_text(nom_alumne, True)
    return [], tokens


def docx_a_pdf_bytes(docx_bytes):
    """Converteix un DOCX a PDF per poder-lo unir amb la resta d'informes.

    Manté el text, els paràgrafs i les taules bàsiques. Les imatges, capçaleres, peus
    i alguns formats avançats poden no conservar-se exactament.
    """
    document = Document(io.BytesIO(docx_bytes))
    buffer_pdf = io.BytesIO()

    pdf = SimpleDocTemplate(
        buffer_pdf,
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
    )

    styles = getSampleStyleSheet()
    estil_normal = styles["BodyText"]
    estil_normal.fontName = "Helvetica"
    estil_normal.fontSize = 10
    estil_normal.leading = 13

    elements = []

    for paragraf in document.paragraphs:
        text = paragraf.text.strip()
        if text:
            elements.append(Paragraph(escape(text), estil_normal))
            elements.append(Spacer(1, 0.15 * cm))

    for taula_docx in document.tables:
        dades = []
        for fila in taula_docx.rows:
            dades.append([Paragraph(escape(cella.text.strip()), estil_normal) for cella in fila.cells])
        if dades:
            taula = Table(dades, repeatRows=1)
            taula.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            elements.append(Spacer(1, 0.2 * cm))
            elements.append(taula)
            elements.append(Spacer(1, 0.3 * cm))

    if not elements:
        elements.append(Paragraph("Document DOCX sense text detectable.", estil_normal))

    pdf.build(elements)
    buffer_pdf.seek(0)
    return buffer_pdf.getvalue()


def prepara_fitxers(files):
    resultat = []
    for f in files or []:
        contingut_original = f.read()
        nom_original = f.name
        nom_sense_extensio = nom_original.rsplit(".", 1)[0]
        extensio = nom_original.rsplit(".", 1)[-1].lower() if "." in nom_original else ""
        avis_conversio = ""

        if extensio == "docx":
            try:
                contingut = docx_a_pdf_bytes(contingut_original)
                avis_conversio = (
                    "Convertit des de DOCX. Revisa el PDF final: el text i les taules bàsiques es conserven, "
                    "però imatges, capçaleres, peus i formats avançats poden variar."
                )
            except Exception as e:
                resultat.append({
                    "nom_original": nom_original,
                    "nom_normalitzat": normalitza_text(nom_sense_extensio),
                    "tokens": set(tokens_text(nom_sense_extensio, True)),
                    "bytes": b"",
                    "error_preparacio": f"No s'ha pogut convertir el DOCX a PDF: {e}",
                    "avis_conversio": "",
                })
                continue
        else:
            contingut = contingut_original

        resultat.append({
            "nom_original": nom_original,
            "nom_normalitzat": normalitza_text(nom_sense_extensio),
            "tokens": set(tokens_text(nom_sense_extensio, True)),
            "bytes": contingut,
            "error_preparacio": "",
            "avis_conversio": avis_conversio,
        })
    return resultat


def puntuacio_coincidencia(nom_alumne, fitxer):
    """
    Retorna una puntuació interna de coincidència.
    No s'ensenya a l'usuari. Serveix per acceptar casos com:
    - 'El Ouahabi Aceituno, Tarek'
    - PDF: 'Tarek El Ouahabi PSI.pdf'
    encara que falti el segon cognom.
    """
    tokens_alumne = set(tokens_text(nom_alumne, True))
    tokens_fitxer = fitxer["tokens"]

    if not tokens_alumne or not tokens_fitxer:
        return 0

    nom_tokens, cognom_tokens = separa_nom_cognoms(nom_alumne)
    nom_tokens = set(nom_tokens)
    cognom_tokens = set(cognom_tokens)

    presents = tokens_alumne & tokens_fitxer
    cobertura = len(presents) / len(tokens_alumne)

    # Cas ideal: totes les parts del nom de l'alumne apareixen al fitxer.
    if tokens_alumne.issubset(tokens_fitxer):
        return 100

    # Format habitual de llistat: cognoms, nom.
    # Accepta nom + algun cognom encara que falti un segon cognom.
    if nom_tokens and cognom_tokens:
        te_nom = bool(nom_tokens & tokens_fitxer)
        cognoms_presents = len(cognom_tokens & tokens_fitxer)
        if te_nom and cognoms_presents >= 1 and len(presents) >= 2:
            return int(80 + min(19, cobertura * 19))

        # També accepta els dos cognoms encara que no aparegui el nom.
        if len(cognom_tokens & tokens_fitxer) >= min(2, len(cognom_tokens)) and len(presents) >= 2:
            return int(70 + min(20, cobertura * 20))

    # Cas general: si coincideixen gairebé totes les parts i com a mínim 2 paraules.
    if len(presents) >= 2 and cobertura >= 0.60:
        return int(60 + min(20, cobertura * 20))

    return 0


def busca_pdf_per_nom(nom_alumne, fitxers):
    """
    Relaciona un PDF amb un alumne segons les parts del nom que apareixen al nom de l'arxiu.
    Prioritza la coincidència completa, però permet que falti algun cognom si hi ha prou seguretat.
    """
    if not fitxers:
        return None, ""

    candidats = []
    for fitxer in fitxers:
        puntuacio = puntuacio_coincidencia(nom_alumne, fitxer)
        if puntuacio > 0:
            tokens_extres = len(fitxer["tokens"] - set(tokens_text(nom_alumne, True)))
            llargada_nom = len(fitxer["nom_normalitzat"])
            candidats.append((puntuacio, -tokens_extres, -llargada_nom, fitxer))

    if not candidats:
        return None, ""

    candidats.sort(key=lambda item: (item[0], item[1], item[2], item[3]["nom_original"]), reverse=True)
    millor = candidats[0][3]

    # Si hi ha més d'una coincidència amb la mateixa força, ho avisem dins del ZIP.
    empatats = [c[3]["nom_original"] for c in candidats[1:] if c[0] == candidats[0][0]]
    if empatats:
        avis = (
            f"Més d'una coincidència possible. S'ha triat {millor['nom_original']}. "
            f"Altres candidats: {', '.join(empatats)}"
        )
        return millor, avis

    return millor, ""


def afegeix_pdf(writer, pdf_bytes):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    for page in reader.pages:
        writer.add_page(page)


def construeix_nom_pdf(alumne, titol_base, format_nom_sortida):
    if format_nom_sortida == "Alumne - Document":
        return nom_seguretat_fitxer(f"{alumne} - {titol_base}.pdf")
    if format_nom_sortida == "Document - Alumne":
        return nom_seguretat_fitxer(f"{titol_base} - {alumne}.pdf")
    return nom_seguretat_fitxer(f"{alumne}.pdf")


def genera_zip(taula, ordre, titol_base, format_nom_sortida):
    buffer_zip = io.BytesIO()
    errors = []
    avisos = []
    generats = 0

    with zipfile.ZipFile(buffer_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for registre in taula:
            writer = PdfWriter()
            afegit = False

            for bloc in ordre:
                fitxer = registre.get(bloc)
                if fitxer is None:
                    continue
                try:
                    if fitxer.get("error_preparacio"):
                        raise ValueError(fitxer["error_preparacio"])
                    afegeix_pdf(writer, fitxer["bytes"])
                    afegit = True
                    if fitxer.get("avis_conversio"):
                        avisos.append(
                            f"{registre['alumne']} - {BLOCS.get(bloc, bloc)} ({fitxer['nom_original']}): "
                            f"{fitxer['avis_conversio']}"
                        )
                except Exception as e:
                    errors.append(
                        f"{registre['alumne']}: error llegint {BLOCS.get(bloc, bloc)} "
                        f"({fitxer['nom_original']}): {e}"
                    )

            if afegit:
                pdf_buffer = io.BytesIO()
                writer.write(pdf_buffer)
                nom_final = construeix_nom_pdf(registre["alumne"], titol_base, format_nom_sortida)
                zf.writestr(nom_final, pdf_buffer.getvalue())
                generats += 1
            else:
                errors.append(f"{registre['alumne']}: no s'ha trobat cap PDF per unificar.")

        for registre in taula:
            for bloc in BLOCS:
                avis = registre.get(f"avis_{bloc}")
                if avis:
                    avisos.append(f"{registre['alumne']} - {BLOCS[bloc]}: {avis}")

        contingut_avisos = avisos + errors
        if contingut_avisos:
            zf.writestr("AVISOS.txt", "\n".join(contingut_avisos))

    buffer_zip.seek(0)
    return buffer_zip, generats, contingut_avisos


st.title("📄 InformeFusion")
st.caption("Unifica informes oficials, PSI i observacions en un únic PDF per alumne. Admet PDF i DOCX.")

with st.expander("Funcionament i protecció de dades", expanded=True):
    st.write(
        "Aquesta aplicació processa els arxius durant la sessió i genera un ZIP final. "
        "No desa els PDFs en cap carpeta permanent. Tot i així, cal utilitzar-la només en un entorn autoritzat pel centre."
    )

with st.expander("Com s'han de pujar els informes oficials?", expanded=True):
    st.markdown(
        """
        - Pots pujar arxius **PDF** i també **DOCX**. Els DOCX es convertiran automàticament a PDF abans d'unificar-los.
        - Si l'informe oficial de notes és **un únic arxiu per alumne**, puja aquests arxius només a **Informe oficial - part 1**.
        - Si l'informe oficial surt dividit en **dos arxius per alumne**, separa abans els arxius en dues carpetes al teu ordinador:
          - una carpeta amb totes les **parts 1**;
          - una altra carpeta amb totes les **parts 2**.
        - Després puja cada grup d'arxius al seu apartat corresponent.
        - No barregis la part 1 i la part 2 en el mateix carregador, perquè l'aplicació no pot saber amb seguretat quin arxiu ha d'anar primer.
        - El PSI i les observacions són opcionals: només s'afegiran als alumnes que tinguin un arxiu corresponent.
        - Si puges DOCX, revisa el resultat final: es conserva el text i les taules bàsiques, però alguns formats avançats poden variar.
        """
    )

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. Llistat d'alumnes")
    excel_alumnes = st.file_uploader(
        "Puja un Excel o CSV amb el llistat de la classe",
        type=["xlsx", "xls", "csv"],
    )

    sense_capcalera = st.checkbox(
        "El llistat no té capçalera i la primera fila ja és un alumne",
        value=False,
    )

    df = None
    alumnes = []
    if excel_alumnes:
        try:
            df = llegeix_alumnes(excel_alumnes, sense_capcalera)
            if sense_capcalera:
                df.columns = [f"Columna {i + 1}" for i in range(len(df.columns))]

            st.success("Llistat carregat correctament.")
            st.dataframe(df.head(10), use_container_width=True)

            columnes = list(df.columns)
            mode_nom = st.radio(
                "Com està escrit el nom de l'alumne?",
                ["Una columna amb nom complet", "Una columna de nom i una de cognoms"],
            )

            if mode_nom == "Una columna amb nom complet":
                columna_nom_complet = st.selectbox("Columna amb el nom complet", columnes)
                columna_nom = columna_cognoms = None
            else:
                columna_nom = st.selectbox("Columna del nom", columnes)
                columna_cognoms = st.selectbox("Columna dels cognoms", columnes)
                columna_nom_complet = None

            alumnes = []
            for _, row in df.iterrows():
                alumne = extreu_nom_alumne(
                    row,
                    mode_nom,
                    columna_nom,
                    columna_cognoms,
                    columna_nom_complet,
                )
                if alumne and alumne.lower() != "nan":
                    alumnes.append(alumne)

            st.info(f"Alumnes detectats: {len(alumnes)}")
        except Exception as e:
            st.error(f"No he pogut llegir el llistat: {e}")

with col2:
    st.subheader("2. Informes")
    pdf_informe_1 = st.file_uploader(
        "Informe oficial - part 1 / notes / creuetes",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        help="Sempre s'ha d'omplir aquest apartat. Si l'informe oficial només té un arxiu per alumne, puja'l aquí. Admet PDF i DOCX.",
    )
    pdf_informe_2 = st.file_uploader(
        "Informe oficial - part 2, només si existeix",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        help="Només s'ha d'omplir si l'informe oficial de notes surt dividit en dos arxius per alumne. Admet PDF i DOCX.",
    )
    pdf_psi = st.file_uploader(
        "PSI / Pla de suport individualitzat, només si n'hi ha",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        help="Es detecta pel nom de l'arxiu. Pot faltar algun segon cognom si el nom i un cognom coincideixen.",
    )
    pdf_observacions = st.file_uploader(
        "Observacions escrites",
        type=["pdf", "docx"],
        accept_multiple_files=True,
    )

st.subheader("3. Configuració")
conf1, conf2 = st.columns([1, 1])
with conf1:
    titol_base = st.text_input("Nom base dels PDFs finals", value="Informe final")
    nom_zip_base = st.text_input("Nom del ZIP final", value="informes_unificats")
with conf2:
    opcions_ordre = list(BLOCS.keys())
    ordre = st.multiselect(
        "Ordre d'unificació",
        options=opcions_ordre,
        default=ORDRE_PER_DEFECTE,
        format_func=lambda codi: BLOCS[codi],
    )
    format_nom_sortida = st.radio(
        "Format del nom dels PDFs finals",
        ["Alumne - Document", "Document - Alumne", "Només alumne"],
        index=0,
        horizontal=True,
    )

st.info(
    "La coincidència es fa pel nom de l'arxiu, sigui PDF o DOCX. L'aplicació prioritza les coincidències completes, "
    "però també accepta casos habituals en què apareixen el nom i un cognom encara que falti un segon cognom."
)

st.subheader("4. Revisió i generació")

if alumnes and ordre:
    fitxers = {
        "informe_1": prepara_fitxers(pdf_informe_1),
        "informe_2": prepara_fitxers(pdf_informe_2),
        "psi": prepara_fitxers(pdf_psi),
        "observacions": prepara_fitxers(pdf_observacions),
    }

    registres = []
    files_revisio = []
    for alumne in alumnes:
        registre = {"alumne": alumne}
        fila = {"Alumne": alumne}

        for bloc, etiqueta in BLOCS.items():
            pdf, avis = busca_pdf_per_nom(alumne, fitxers[bloc])
            registre[bloc] = pdf
            registre[f"avis_{bloc}"] = avis
            fila[etiqueta] = pdf["nom_original"] if pdf else "—"

        registres.append(registre)
        files_revisio.append(fila)

    st.dataframe(pd.DataFrame(files_revisio), use_container_width=True)

    if st.button("Generar ZIP amb els informes unificats", type="primary"):
        zip_buffer, generats, avisos = genera_zip(registres, ordre, titol_base, format_nom_sortida)
        data_actual = datetime.now().strftime("%Y%m%d_%H%M")
        st.success(f"PDFs generats: {generats}")
        if avisos:
            st.warning("Hi ha avisos. També trobaràs un fitxer AVISOS.txt dins del ZIP.")
        st.download_button(
            "Descarregar ZIP",
            data=zip_buffer,
            file_name=nom_seguretat_fitxer(f"{nom_zip_base}_{data_actual}.zip"),
            mime="application/zip",
        )
else:
    st.info("Puja el llistat d'alumnes i selecciona l'ordre per començar.")
